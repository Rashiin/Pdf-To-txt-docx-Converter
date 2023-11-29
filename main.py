import PyPDF2
import docx
import persian
from langdetect import detect
from tkinter import Tk, Text, messagebox, filedialog, Label
from tkinter.ttk import Button, Style


class PDFToTextConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF to Text Converter")
        self.create_widgets()
        self.style_widgets()


    def create_widgets(self):
        self.upload_btn = Button(self.root, text="📂 Upload PDF", command=self.upload_action)
        self.upload_btn.pack(pady=10)


        self.page_num_label = Label(self.root, text="Number of pages: ")
        self.page_num_label.pack()


        self.text_output = Text(self.root, height=10, width=50, wrap='word')
        self.text_output.pack(padx=10, pady=10)


        self.save_text_btn = Button(self.root, text="💾 Save as Text", command=self.save_text_action)
        self.save_text_btn.pack(pady=10)


        self.save_word_btn = Button(self.root, text="💾 Save as Word", command=self.save_word_action)
        self.save_word_btn.pack(pady=10)


    def style_widgets(self):
        style = Style()
        style.theme_use('clam')
        style.configure('TButton', background='#ff80ed', foreground='white',
                        font=('Helvetica', 11), width=20, borderwidth=1)
        style.map('TButton', background=[('active', '#ff4ddb')])


        style.configure('TLabel', background='#f0a3ff', foreground='black',
                        font=('Helvetica', 11))


        self.root.configure(background='#000000')
        self.text_output.configure(font=('Helvetica', 11),
                                   selectbackground='yellow', bd=0, relief='flat')


    def upload_action(self):
        file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if file_path:
            try:
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    self.page_num_label.config(text=f"Number of pages: {len(pdf_reader.pages)}")


                    pdf_obj = pdf_reader.pages[0]
                    my_text = pdf_obj.extract_text()


                    self.lang = detect(my_text)


                    self.text_output.delete('1.0', 'end')
                    self.text_output.insert('end', my_text)
            except Exception as e:
                messagebox.showerror("Error", f"❌ An error occurred: {e}")


    def save_text_action(self):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
        )


        if file_path:
            try:
                with open(file_path, "w", encoding='utf-8') as file:
                    text = self.text_output.get("1.0", 'end')
                    file.writelines(text)
                messagebox.showinfo("Success", "✅ The text has been saved!")
            except Exception as e:
                messagebox.showerror("Error", f"❌ An error occurred: {e}")


    def save_word_action(self):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )


        if file_path:
            try:
                doc = docx.Document()
                text = self.text_output.get("1.0", 'end')


                if self.lang == 'fa':
                    text = persian.convert_ar_characters(text)
                    paragraph = doc.add_paragraph(text)
                    paragraph.paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT
                    for run in paragraph.runs:
                        run.font.rtl = True
                else:
                    doc.add_paragraph(text)


                doc.save(file_path)
                messagebox.showinfo("Success", "✅ The Word document has been saved!")
            except Exception as e:
                messagebox.showerror("Error", f"❌ An error occurred: {e}")


def main():
    root = Tk()
    app = PDFToTextConverter(root)
    root.mainloop()


if __name__ == "__main__":
    main()



""" Developed By Rashin Gholijani Farahani """